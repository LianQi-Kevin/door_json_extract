from pathlib import Path

from docx import Document


def iter_unique_paragraphs(doc):
    seen = set()

    def walk(_container):
        for p in _container.paragraphs:
            if p._p not in seen:
                seen.add(p._p)
                yield p

        for table in _container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from walk(cell)

    yield from walk(doc)

    for section in doc.sections:
        for container in [section.header, section.footer, section.first_page_header, section.first_page_footer,
                          section.even_page_header, section.even_page_footer, ]:
            yield from walk(container)


def replace_token_in_paragraph(paragraph, token, replacement):
    full_text = "".join(run.text for run in paragraph.runs)
    start = full_text.index(token)
    end = start + len(token)

    cursor = 0
    replaced = False

    for run in paragraph.runs:
        text = run.text
        run_start = cursor
        run_end = cursor + len(text)
        cursor = run_end

        if run_end <= start or run_start >= end:
            continue

        overlap_start = max(start, run_start)
        overlap_end = min(end, run_end)

        prefix = text[: overlap_start - run_start]
        suffix = text[overlap_end - run_start:]

        if not replaced:
            run.text = prefix + replacement + suffix
            replaced = True
        else:
            run.text = suffix


def replace_once(doc, token, replacement):
    matches = [p for p in iter_unique_paragraphs(doc) if token in "".join(run.text for run in p.runs)]

    if len(matches) != 1:
        raise ValueError(f"占位符 {token} 匹配数量不是 1，而是 {len(matches)}")

    replace_token_in_paragraph(matches[0], token, replacement)


def fill_upgrade_request_docx(template_path, output_path, drawing_name_title, drawing_names):
    doc = Document(template_path)

    replace_once(doc, "<图纸名称标题>", drawing_name_title)
    replace_once(doc, "<图纸名称>", "\n".join(drawing_names))

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
